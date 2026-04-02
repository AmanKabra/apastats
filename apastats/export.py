"""
APA 7th edition table export to Word (.docx) and CSV.

Word tables follow APA formatting:
  - Times New Roman, 12pt body, 10pt notes
  - No vertical lines
  - Horizontal borders: top of table, below header, bottom of table
  - Table title: bold number, italic title
  - Notes below the bottom border

References
----------
APA (2020). *Publication Manual of the APA* (7th ed.), Chapter 7.
"""

from __future__ import annotations

import csv
import io
from pathlib import Path
from typing import Optional, Sequence, Union

import pandas as pd


# ═══════════════════════════════════════════════════════════════════════════
# Word (.docx) export
# ═══════════════════════════════════════════════════════════════════════════

def to_docx(
    table_df: pd.DataFrame,
    filepath: Union[str, Path],
    title: str = "Table 1",
    subtitle: str = "",
    note: str = "",
    font_name: str = "Times New Roman",
    font_size_body: int = 11,
    font_size_note: int = 10,
) -> Path:
    """Export a DataFrame as an APA-formatted Word table.

    Parameters
    ----------
    table_df : pd.DataFrame
        The formatted table (e.g. from ``DescriptivesResult.table_df``).
    filepath : str or Path
        Output ``.docx`` path.
    title : str
        Table number line (e.g. ``"Table 1"``).  Rendered bold.
    subtitle : str
        Table title line (e.g. ``"Means, Standard Deviations, …"``).
        Rendered italic.
    note : str
        Table note text (rendered below the table in smaller font).
    font_name : str
        Font family (default ``"Times New Roman"``).
    font_size_body : int
        Font size in points for table content (default 11).
    font_size_note : int
        Font size in points for notes (default 10).

    Returns
    -------
    Path
        The written file path.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # --- Page setup: 1-inch margins ---
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- Title: bold table number ---
    p_title = doc.add_paragraph()
    run = p_title.add_run(title)
    run.bold = True
    run.font.name = font_name
    run.font.size = Pt(font_size_body)
    p_title.paragraph_format.space_after = Pt(0)

    # --- Subtitle: italic ---
    if subtitle:
        p_sub = doc.add_paragraph()
        run = p_sub.add_run(subtitle)
        run.italic = True
        run.font.name = font_name
        run.font.size = Pt(font_size_body)
        p_sub.paragraph_format.space_after = Pt(2)

    # --- Table ---
    n_rows = len(table_df)
    n_cols = len(table_df.columns)
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Table Grid"  # we'll remove vertical borders

    # Helper: set cell text with font
    def _set_cell(cell, text: str, bold: bool = False, align_right: bool = False):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.name = font_name
        run.font.size = Pt(font_size_body)
        run.bold = bold
        if align_right:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Remove cell margins for tighter layout
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement("w:tcMar")
        for side in ("top", "bottom", "start", "end"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), "40")
            el.set(qn("w:type"), "dxa")
            tcMar.append(el)
        tcPr.append(tcMar)

    # Header row
    for j, col_name in enumerate(table_df.columns):
        _set_cell(
            table.rows[0].cells[j], col_name,
            bold=True,
            align_right=(j > 0),  # first column left-aligned
        )

    # Data rows
    for i in range(n_rows):
        for j, col_name in enumerate(table_df.columns):
            val = table_df.iloc[i, j]
            _set_cell(
                table.rows[i + 1].cells[j], str(val),
                align_right=(j > 0),
            )

    # --- APA border formatting: remove all borders, then add only ---
    # top border on first row, bottom border on header, bottom border on last row
    def _set_border(cell, sides: dict):
        """Set specific borders on a cell.
        sides: dict of side -> {'val': 'single', 'sz': '8', 'color': '000000'}
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side, attrs in sides.items():
            el = OxmlElement(f"w:{side}")
            for k, v in attrs.items():
                el.set(qn(f"w:{k}"), v)
            tcBorders.append(el)
        tcPr.append(tcBorders)

    border_on = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    border_off = {"val": "none", "sz": "0", "space": "0", "color": "auto"}

    for i, row in enumerate(table.rows):
        for cell in row.cells:
            sides = {
                "top": border_off,
                "bottom": border_off,
                "start": border_off,  # no vertical lines
                "end": border_off,
            }
            # Top of table (row 0, top)
            if i == 0:
                sides["top"] = border_on
            # Below header (row 0, bottom)
            if i == 0:
                sides["bottom"] = border_on
            # Bottom of table (last row, bottom)
            if i == len(table.rows) - 1:
                sides["bottom"] = border_on
            _set_border(cell, sides)

    # --- Note ---
    if note:
        p_note = doc.add_paragraph()
        # "Note." in italic
        if note.startswith("Note."):
            run_note_label = p_note.add_run("Note. ")
            run_note_label.italic = True
            run_note_label.font.name = font_name
            run_note_label.font.size = Pt(font_size_note)
            remaining = note[len("Note. "):]
            run_rest = p_note.add_run(remaining)
            run_rest.font.name = font_name
            run_rest.font.size = Pt(font_size_note)
        else:
            run = p_note.add_run(note)
            run.font.name = font_name
            run.font.size = Pt(font_size_note)

    filepath = Path(filepath)
    doc.save(filepath)
    return filepath



# ═══════════════════════════════════════════════════════════════════════════
# CSV export
# ═══════════════════════════════════════════════════════════════════════════

def to_csv(
    table_df: pd.DataFrame,
    filepath: Optional[Union[str, Path]] = None,
) -> str:
    """Export a DataFrame to CSV.

    Parameters
    ----------
    table_df : pd.DataFrame
    filepath : str or Path, optional
        If provided, writes to file and returns the path.
        If None, returns the CSV string.

    Returns
    -------
    str
        CSV content (or filepath as string if written to disk).
    """
    buf = io.StringIO()
    table_df.to_csv(buf, index=False, quoting=csv.QUOTE_NONNUMERIC)
    content = buf.getvalue()

    if filepath is not None:
        filepath = Path(filepath)
        filepath.write_text(content, encoding="utf-8")
        return str(filepath)
    return content

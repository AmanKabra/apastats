"""Tests for APA export to Word, LaTeX, and CSV."""

import os
import tempfile
import pandas as pd
import pytest
from apastats.export import to_docx, to_latex, to_csv


@pytest.fixture
def sample_table():
    """A simple formatted table DataFrame."""
    return pd.DataFrame({
        "Variable": ["1. POS", "2. Commitment", "3. Performance"],
        "M": ["3.48", "3.80", "3.08"],
        "SD": ["0.85", "0.92", "0.75"],
        "1": ["(.99)", "", ""],
        "2": [".50**", "(.98)", ""],
        "3": [".28**", ".40**", "(.97)"],
    })


class TestDocxExport:
    def test_creates_file(self, sample_table):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            result = to_docx(
                sample_table, path,
                title="Table 1",
                subtitle="Means, Standard Deviations, and Intercorrelations",
                note="Note. N = 400. *p < .05. **p < .01.",
            )
            assert os.path.exists(result)
            assert os.path.getsize(result) > 0
        finally:
            os.unlink(path)

    def test_no_crash_empty_note(self, sample_table):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            to_docx(sample_table, path, title="Table 1")
            assert os.path.exists(path)
        finally:
            os.unlink(path)

    def test_docx_content(self, sample_table):
        """Verify the .docx contains table data."""
        from docx import Document
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            to_docx(sample_table, path, title="Table 1")
            doc = Document(path)
            # Should have at least one table
            assert len(doc.tables) == 1
            table = doc.tables[0]
            # Header row + 3 data rows
            assert len(table.rows) == 4
            # First data cell should be "1. POS"
            assert "POS" in table.rows[1].cells[0].text
        finally:
            os.unlink(path)


class TestLatexExport:
    def test_basic_output(self, sample_table):
        latex = to_latex(sample_table, title="Table 1",
                         subtitle="Descriptive Statistics")
        assert "\\begin{table}" in latex
        assert "\\toprule" in latex
        assert "\\midrule" in latex
        assert "\\bottomrule" in latex
        assert "\\end{table}" in latex

    def test_contains_data(self, sample_table):
        latex = to_latex(sample_table)
        assert "POS" in latex
        assert "3.48" in latex

    def test_column_alignment(self, sample_table):
        latex = to_latex(sample_table)
        # First column l, rest r
        assert "lrrrrr" in latex

    def test_note(self, sample_table):
        latex = to_latex(sample_table, note="Note. N = 400.")
        assert "footnotesize" in latex


class TestCsvExport:
    def test_returns_string(self, sample_table):
        result = to_csv(sample_table)
        assert isinstance(result, str)
        assert "Variable" in result
        assert "POS" in result

    def test_writes_file(self, sample_table):
        with tempfile.NamedTemporaryFile(suffix=".csv", delete=False, mode="w") as f:
            path = f.name
        try:
            result = to_csv(sample_table, filepath=path)
            assert os.path.exists(path)
            content = open(path, encoding="utf-8").read()
            assert "Variable" in content
        finally:
            os.unlink(path)

    def test_roundtrip(self, sample_table):
        """CSV can be re-read as a DataFrame."""
        import io
        csv_str = to_csv(sample_table)
        df_back = pd.read_csv(io.StringIO(csv_str))
        assert list(df_back.columns) == list(sample_table.columns)
        assert len(df_back) == len(sample_table)
